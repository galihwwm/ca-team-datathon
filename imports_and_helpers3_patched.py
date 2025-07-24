# --- Imports and Shared Helpers ---
import os
import re
import uuid
import pymupdf
import openai
import chromadb
from dotenv import load_dotenv
from llama_index.core import Document, VectorStoreIndex, StorageContext, Settings
from llama_index.vector_stores.chroma import ChromaVectorStore
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.response_synthesizers import CompactAndRefine
from llama_index.llms.openai import OpenAI
from llama_index.core.tools import QueryEngineTool, ToolMetadata
from llama_index.multi_modal_llms.openai import OpenAIMultiModal
from llama_index.core.agent.react_multimodal.step import MultimodalReActAgentWorker
from llama_index.core.agent import Task, AgentRunner
from llama_index.core.schema import ImageDocument
import pymupdf4llm
import pathlib
from docx import Document as DocxDocument
from datetime import datetime
import asyncio
from concurrent.futures import ThreadPoolExecutor
from collections import defaultdict
from glob import glob
from concurrent.futures import ThreadPoolExecutor
from tiktoken import encoding_for_model
from openai import OpenAI

load_dotenv()
openai_client = openai.OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

# --- PATCH for from_openai_message (Avoid recursion) ---
from llama_index.multi_modal_llms.openai import base as openai_mm_base
original_from_openai_message = openai_mm_base.from_openai_message

def patched_from_openai_message(openai_message, modalities=None):
    if modalities is None:
        modalities = ["text", "image"]
    return original_from_openai_message(openai_message, modalities)

openai_mm_base.from_openai_message = patched_from_openai_message

# --- Filename Helpers ---
def sanitize_filename(name):
    name = name.strip().replace(" ", "_").replace("/", "_").replace("\\", "_")
    name = re.sub(r"[^a-zA-Z0-9_\.-]", "", name)
    return name[:50]

def normalize_text(text):
    return re.sub(r'[^a-z0-9]', '', text.lower())

# --- Save Output as .docx ---
def save_to_docx(content, prefix="report"):
    doc = DocxDocument()
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    doc.add_heading("Evaluation/Guidance Report", level=1)
    doc.add_paragraph(content)
    filepath = f"{prefix}_{timestamp}.docx"
    doc.save(filepath)
    print(f"📄 Saved report to {filepath}")

# --- Report Generator Class ---
from docx.shared import Pt
def bold_keywords(paragraph, keywords=["Compliant", "Not Compliant", "Partially Compliant", "Not Applicable"]):
    for keyword in keywords:
        if keyword in paragraph.text:
            parts = paragraph.text.split(keyword)
            paragraph.clear()
            paragraph.add_run(parts[0])
            run = paragraph.add_run(keyword)
            run.bold = True
            run.font.size = Pt(11)
            if len(parts) > 1:
                paragraph.add_run(parts[1])
            break

class ReportGenerator:
    def __init__(self, prefix="report"):
        self.prefix = prefix
        self.doc = DocxDocument()
        self.timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        self.doc.add_heading("Compliance Evaluation Report", 0)

    def add_section(self, title, content):
        self.doc.add_heading(title, level=2)
        self.doc.add_paragraph(content)

    
    def add_results(self, results, result_type="Evaluation"):
        for res in results:
            if result_type == "Evaluation":
                title = res.get("workunit", "General Evaluation Result")
                body = res.get("evaluation", res.get("response", ""))
                self.doc.add_heading(f"📋 Evaluation Result: {title}", level=3)
                self.doc.add_paragraph(body)
                bold_keywords(self.doc.paragraphs[-1])
            elif result_type == "Developer":
                title = res.get("developer_action", "General Developer Guidance")
                body = res.get("guidance", res.get("response", ""))
                #evidence_excerpt = res.get("evidence_excerpt_used", "")
                self.doc.add_heading(f"🛠 Developer Guidance: {title}", level=3)
                # Add the evidence excerpt first
                # if evidence_excerpt:
                #     self.doc.add_paragraph("📄 *Evidence excerpt used in response:*")
                #     self.doc.add_paragraph(evidence_excerpt[:2000] + "..." if len(evidence_excerpt) > 2000 else evidence_excerpt)
                self.doc.add_paragraph(body)
                bold_keywords(self.doc.paragraphs[-1])
    def add_background(self, background_knowledge):
            self.doc.add_heading("📚 Background Knowledge", level=2)
            for part, text in background_knowledge.items():
                self.doc.add_heading(f"Part {part}", level=3)
                self.doc.add_paragraph(str(text).strip())  # Convert Response to string before strip

    def save(self):
        filename = f"{self.prefix}_{self.timestamp}.docx"
        os.makedirs("outputs", exist_ok=True)
        filepath = os.path.join("outputs", filename)
        self.doc.save(filepath)
        print(f"📄 Report saved to {filepath}")
        return filepath

# --- Refactored Markdown Conversion with Logging ---
def parse_pdf_to_markdown_with_images(filepath):
    doc = pymupdf.open(filepath)
    image_dir = "images"
    os.makedirs(image_dir, exist_ok=True)
    md_chunks = []
    image_documents = []
    image_index = 1

    all_image_titles = []

    for page_num, page in enumerate(doc):
        blocks = page.get_text("dict")["blocks"]
        page_md = ""
        img_names = []
        surrounding_texts = []

        figure_titles = {}
        for block in blocks:
            if block.get("type") == 0:
                for line in block.get("lines", []):
                    text = " ".join(span["text"] for span in line["spans"] if span["text"].strip())
                    page_md += text + "\n"
                    surrounding_texts.append(text)
                    match = re.match(r"Figure\s*\d+[:\.]?\s*(.*)", text, re.IGNORECASE)
                    if match:
                        figure_titles[len(surrounding_texts)] = match.group(1).strip()

        for block in blocks:
            if block.get("type") == 1:
                img_list = page.get_images(full=True)
                for img in img_list:
                    xref = img[0]
                    base_image = doc.extract_image(xref)
                    ext = base_image["ext"]
                    img_bytes = base_image["image"]

                    closest_title = None
                    for idx in sorted(figure_titles.keys(), reverse=True):
                        if idx <= len(surrounding_texts):
                            closest_title = figure_titles[idx]
                            break

                    if not closest_title:
                        context_snippet = " ".join(surrounding_texts[-3:])[:150]
                        closest_title = context_snippet.split(".")[0] or f"Image_{image_index}"

                    safe_title = sanitize_filename(closest_title)
                    img_name = f"{safe_title}_p{page_num+1}_{image_index}.{ext}"
                    img_path = os.path.join(image_dir, img_name)

                    with open(img_path, "wb") as f:
                        f.write(img_bytes)

                    page_md += f"\n![{closest_title}]({img_path})\n"
                    img_names.append(img_name)
                    all_image_titles.append(closest_title)
                    image_documents.append(ImageDocument(image_path=img_path, text=closest_title))
                    image_index += 1

        md_chunks.append(Document(
            text=page_md.strip(),
            metadata={
                "source": filepath,
                "page": page_num + 1,
                "image_count": len(img_names),
                "image_titles": ", ".join(img_names)
            }
        ))

    # markdown_preview_path = f"preview_markdown_{sanitize_filename(os.path.basename(filepath))}.txt"
    # with open(markdown_preview_path, "w", encoding="utf-8") as f:
    #     for doc in md_chunks:
    #         f.write(doc.text + "\n\n")

    # image_titles_path = f"image_titles_{sanitize_filename(os.path.basename(filepath))}.txt"
    # with open(image_titles_path, "w", encoding="utf-8") as f:
    #     for title in sorted(set(all_image_titles)):
    #         f.write(title + "\n")

    return md_chunks, image_documents

def make_image_document(page, page_number, doc, img_index, image_output_dir="images"):
    xref = page.get_images(full=True)[img_index][0]
    base_image = doc.extract_image(xref)
    image_bytes = base_image["image"]
    image_ext = base_image["ext"]

    image_path = os.path.join(image_output_dir, f"page{page_number}_img{img_index}.{image_ext}")
    with open(image_path, "wb") as f:
        f.write(image_bytes)

    # --- Try to infer surrounding text ---
    blocks = page.get_text("dict")["blocks"]
    nearby_text = []
    for block in blocks:
        if block.get("type") == 0:  # text block
            for line in block.get("lines", []):
                line_text = " ".join(span["text"] for span in line["spans"] if span["text"].strip())
                if line_text:
                    nearby_text.append(line_text)

    context_text = " ".join(nearby_text[-5:])  # last few lines (often caption is after image)
    has_caption = bool(context_text.strip())

    caption = context_text[:120] if has_caption else f"Uncaptioned image from page {page_number}"
    print(f"🖼️ Saved image: {image_path} | Caption: {caption}")

    return ImageDocument(
        image_path=image_path,
        text=caption,
        metadata={
            "page": page_number,
            "caption": caption,
            "surrounding_text": context_text,
            "image_path": image_path,
            "image_ext": image_ext,
            "modality": "image"
        }
    )



def parse_pdf_to_indexable_documents(pdf_path, image_output_dir="images"):
    os.makedirs(image_output_dir, exist_ok=True)

    doc = pymupdf.open(pdf_path)
    indexable_docs = []

    for page_number, page in enumerate(doc, start=1):
        # --- Extract text per page ---
        blocks = page.get_text("dict")["blocks"]
        page_text = []
        for block in blocks:
            if "lines" in block:
                for line in block["lines"]:
                    text = " ".join([span["text"] for span in line["spans"] if span["text"].strip()])
                    if text:
                        page_text.append(text)
        if page_text:
            indexable_docs.append(Document(
                text="\n".join(page_text),
                metadata={"page": page_number, "modality": "text"}
            ))
        print(f"✅ Indexed page {page_number}: {len(page_text)} text lines")

        # --- Extract images per page ---
        img_list = page.get_images(full=True)
        for img_index, _ in enumerate(img_list):
            image_doc = make_image_document(page, page_number, doc, img_index, image_output_dir)
            indexable_docs.append(image_doc)
            print(f"✅ Indexed image → Page {page_number} | Path: {image_doc.metadata.get('image_path')} | Caption: {image_doc.metadata.get('caption')}")

    return indexable_docs